"""知识库管理路由：文档列出 / 上传（PDF·Word·文本）/ 删除（管理员权限）。"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from pydantic import BaseModel

from .. import kb as kb_mod
from ..deps import require_role

router = APIRouter(prefix="/api/kb", tags=["kb"])

# 上传文件大小上限（字节）
MAX_UPLOAD_BYTES = 10_000_000

# 支持的文件类型 -> 显示名
SUPPORTED_EXTS = {
    ".txt": "文本",
    ".md":  "Markdown",
    ".pdf": "PDF",
    ".docx": "Word",
}


class DocOut(BaseModel):
    id: int
    filename: str
    title: str
    tag: str
    content: str
    created_at: str


def _serialize(d: dict) -> DocOut:
    return DocOut(id=d["id"], filename=d["filename"], title=d["title"],
                  tag=d["tag"], content=d["content"], created_at=d["created_at"])


def _ext_of(name: str) -> str:
    return ("." + name.rsplit(".", 1)[-1]).lower() if "." in name else ""


def _extract_text(filename: str, data: bytes) -> str:
    """按扩展名解析文件为纯文本；不支持的类型返回 None 由调用方报错。"""
    ext = _ext_of(filename)
    if ext in (".txt", ".md"):
        return data.decode("utf-8", errors="replace").strip()
    if ext == ".pdf":
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()
    if ext == ".docx":
        from docx import Document
        import io
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        # 也提取表格单元格文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts).strip()
    return ""


@router.get("", response_model=list[DocOut])
def list_docs(ctx: dict = Depends(require_role("operator", "admin"))):
    return [_serialize(d) for d in kb_mod.list_documents()]


@router.post("/upload", response_model=DocOut, status_code=201)
async def upload(file: UploadFile = File(...),
                 ctx: dict = Depends(require_role("operator", "admin"))):
    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="文件过大（上限 10MB）")
    filename = file.filename or "document.txt"
    ext = _ext_of(filename)
    if ext not in SUPPORTED_EXTS:
        raise HTTPException(status_code=415, detail=f"不支持的文件类型: {ext or '无扩展名'}（支持 {', '.join(SUPPORTED_EXTS)}）")
    try:
        raw = _extract_text(filename, data)
    except Exception as e:
        raise HTTPException(status_code=400,
                            detail=f"文件解析失败: {type(e).__name__}")
    if not raw:
        raise HTTPException(status_code=400, detail="未能从文件提取到文本（可能是扫描件/空文档）")
    # 标题取文件名（去扩展名），标签固定 doc
    title = filename.rsplit(".", 1)[0]
    kb_mod.add_document(filename=filename, title=title, content=raw, tag="doc")
    return _serialize(kb_mod.list_documents()[-1])


@router.delete("/{doc_id}", status_code=200)
def delete(doc_id: int, ctx: dict = Depends(require_role("operator", "admin"))):
    if not kb_mod.delete_document(doc_id):
        raise HTTPException(status_code=404, detail="文档不存在")
    return {"ok": True}
